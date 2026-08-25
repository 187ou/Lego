"""多格式文档加载器

支持格式：
- PDF (.pdf) → 文本提取 + 分页
- 图片 (.png, .jpg, .jpeg, .bmp, .webp) → OCR 文本提取
- 文本 (.txt, .md, .markdown) → 直接读取
- Word (.docx) → 文本提取

统一输出：list[Document]（LangChain Document 格式）
"""

import os
import re
from enum import Enum
from typing import Optional
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


class DocumentType(str, Enum):
    """文档类型枚举"""
    PDF = "pdf"
    IMAGE = "image"
    TEXT = "text"
    DOCX = "docx"
    UNKNOWN = "unknown"


# 文件扩展名到文档类型的映射
EXTENSION_MAP = {
    ".pdf": DocumentType.PDF,
    ".png": DocumentType.IMAGE,
    ".jpg": DocumentType.IMAGE,
    ".jpeg": DocumentType.IMAGE,
    ".bmp": DocumentType.IMAGE,
    ".webp": DocumentType.IMAGE,
    ".txt": DocumentType.TEXT,
    ".md": DocumentType.TEXT,
    ".markdown": DocumentType.TEXT,
    ".docx": DocumentType.DOCX,
}


def detect_document_type(file_path: str) -> DocumentType:
    """根据文件扩展名检测文档类型"""
    ext = os.path.splitext(file_path)[1].lower()
    return EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)


class DocumentLoader:
    """多格式文档加载器"""

    def __init__(self):
        self._parsers = {
            DocumentType.PDF: self._parse_pdf,
            DocumentType.IMAGE: self._parse_image,
            DocumentType.TEXT: self._parse_text,
            DocumentType.DOCX: self._parse_docx,
        }

    def load(
        self,
        file_path: str,
        set_id: str = "",
        metadata: Optional[dict] = None,
    ) -> list[Document]:
        """
        加载文档并返回 Document 列表。

        Args:
            file_path: 文件路径
            set_id: 套装编号
            metadata: 额外元数据

        Returns:
            切片后的 Document 列表

        Raises:
            ValueError: 不支持的文档格式
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        doc_type = detect_document_type(file_path)
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(
                f"不支持的文档格式: {file_path}\n"
                f"支持的格式: {', '.join(EXTENSION_MAP.keys())}"
            )

        parser = self._parsers[doc_type]
        documents = parser(file_path, set_id)

        # 添加额外元数据
        if metadata:
            for doc in documents:
                doc.metadata.update(metadata)

        # 添加文档类型
        for doc in documents:
            doc.metadata["doc_type"] = doc_type.value
            doc.metadata["source_file"] = os.path.basename(file_path)

        return documents

    def _parse_pdf(self, file_path: str, set_id: str) -> list[Document]:
        """解析 PDF 文档"""
        from pypdf import PdfReader
        # RecursiveCharacterTextSplitter 已在文件顶部导入

        reader = PdfReader(file_path)
        documents = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if not text or not text.strip():
                continue

            doc = Document(
                page_content=text.strip(),
                metadata={
                    "set_id": set_id,
                    "page_number": page_num,
                    "source": file_path,
                },
            )
            documents.append(doc)

        # 进一步切片（按段落）
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", ". ", " "],
        )

        split_docs = splitter.split_documents(documents)

        # 重新添加元数据
        for doc in split_docs:
            doc.metadata.setdefault("set_id", set_id)

        return split_docs

    def _parse_image(self, file_path: str, set_id: str) -> list[Document]:
        """
        解析图片文档（OCR 文本提取）。

        优先使用 pytesseract，不可用时使用 PIL 提取基本信息。
        """
        text = self._ocr_image(file_path)

        if not text.strip():
            # OCR 失败时返回基本信息
            text = f"[图片] {os.path.basename(file_path)}"

        return [
            Document(
                page_content=text.strip(),
                metadata={
                    "set_id": set_id,
                    "page_number": 1,
                    "source": file_path,
                    "is_image": True,
                },
            )
        ]

    def _ocr_image(self, image_path: str) -> str:
        """OCR 图片提取文本"""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(image_path)
            # 支持中英文识别
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            return text.strip()
        except ImportError:
            # pytesseract 未安装
            return ""
        except Exception as e:
            print(f"[WARN] OCR 失败: {e}")
            return ""

    def _parse_text(self, file_path: str, set_id: str) -> list[Document]:
        """解析纯文本/Markdown 文档"""
        # RecursiveCharacterTextSplitter 已在文件顶部导入

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        if not content.strip():
            return []

        # 按段落切片
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", "。", ". ", " "],
        )

        # 先创建一个 Document 再切片
        doc = Document(
            page_content=content.strip(),
            metadata={
                "set_id": set_id,
                "source": file_path,
            },
        )

        split_docs = splitter.split_documents([doc])

        for d in split_docs:
            d.metadata.setdefault("set_id", set_id)

        return split_docs

    def _parse_docx(self, file_path: str, set_id: str) -> list[Document]:
        """解析 Word 文档"""
        try:
            from docx import Document as DocxDocument

            docx = DocxDocument(file_path)
            paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            if not content.strip():
                return []

            # RecursiveCharacterTextSplitter 已在文件顶部导入

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50,
                separators=["\n\n", "\n", "。", ". ", " "],
            )

            doc = Document(
                page_content=content.strip(),
                metadata={
                    "set_id": set_id,
                    "source": file_path,
                },
            )

            split_docs = splitter.split_documents([doc])

            for d in split_docs:
                d.metadata.setdefault("set_id", set_id)

            return split_docs

        except ImportError:
            raise ValueError(
                "解析 Word 文档需要安装 python-docx: pip install python-docx"
            )


# 全局单例
_loader: Optional[DocumentLoader] = None


def get_document_loader() -> DocumentLoader:
    """获取文档加载器单例"""
    global _loader
    if _loader is None:
        _loader = DocumentLoader()
    return _loader


def load_document(
    file_path: str,
    set_id: str = "",
    metadata: Optional[dict] = None,
) -> list[Document]:
    """
    便捷函数：加载文档。

    Args:
        file_path: 文件路径
        set_id: 套装编号
        metadata: 额外元数据

    Returns:
        切片后的 Document 列表
    """
    loader = get_document_loader()
    return loader.load(file_path, set_id=set_id, metadata=metadata)
